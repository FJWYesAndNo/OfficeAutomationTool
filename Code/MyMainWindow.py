# -*- coding:GBK -*-

import sys
from PyQt5.QtWidgets import  (QApplication, QMainWindow, QFileDialog)
from PyQt5.QtCore import  pyqtSlot,QDir
import pandas as pd
import os
from docx import Document
import xlrd
from win32com.client import Dispatch,constants,gencache
import pdfplumber
from smtplib import SMTP_SSL
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.header import Header

from ui_MainWindow import Ui_MainWindow
class QmyMainWindow(QMainWindow):

    def __init__(self, parent=None):
        super().__init__(parent)   #���ø��๹�캯������������
        self.ui=Ui_MainWindow()    #����UI����
        self.ui.setupUi(self)      #����UI����

        self.setWindowTitle("Office Automation Tool")
        self.setCentralWidget(self.ui.tabWidget)



#====================Excel���====================
#====================���ϲ�====================
    #���ϲ�����
    def __ExcelMerge(self,excellist='',saveexcel=''):
        # �ļ�·��
        file_dir = excellist
        # �����µı������
        new_filename = saveexcel
        # �ҵ��ļ�·���µ����б�����ƣ������б�
        file_list = os.listdir(file_dir)
        new_list = []

        for file in file_list:
            # �ع��ļ�·��
            file_path = os.path.join(file_dir, file)
            # ��excelת����DataFrame
            dataframe = pd.read_excel(file_path)
            # ���浽���б���
            new_list.append(dataframe)

        # ���DataFrame�ϲ�Ϊһ��
        df = pd.concat(new_list)
        # д�뵽һ����excel����
        df.to_excel(new_filename, index=False)

    @pyqtSlot() #ѡ��Ҫ�ϲ����ļ���
    def on_btn_MergeExcelList_clicked(self):
        curPath=QDir.currentPath()  #��ȡϵͳ��ǰĿ¼
        dlgTitle="ѡ��Ŀ¼"
        selectedDir=QFileDialog.getExistingDirectory(self,
                    dlgTitle,curPath,QFileDialog.ShowDirsOnly)
        self.ui.lineEditMergeExcelList.setText(selectedDir)

    @pyqtSlot() #ѡ�񱣴�·��
    def on_btn_SaveMergeExcel_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '�����ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getSaveFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_SaveMergeExcel.setText(filename)

    @pyqtSlot() #һ���ϲ�
    def on_btn_ExcelMerge_clicked(self):
        try:
            self.ui.textEdit_ExcelStatus.clear()
            excellist=self.ui.lineEditMergeExcelList.text()
            saveexcel=self.ui.lineEdit_SaveMergeExcel.text()
            self.__ExcelMerge(excellist=excellist,saveexcel=saveexcel)
            self.ui.textEdit_ExcelStatus.setPlainText("�ϲ��ɹ���")
            self.ui.lineEditMergeExcelList.clear()
            self.ui.lineEdit_SaveMergeExcel.clear()
        except:
            self.ui.textEdit_ExcelStatus.clear()
            self.ui.textEdit_ExcelStatus.setPlainText("�ϲ�ʧ�ܣ������Ƿ������ȷ��")
            self.ui.lineEditMergeExcelList.clear()
            self.ui.lineEdit_SaveMergeExcel.clear()

#====================�����ظ�====================
    #�����ظ�����
    def __ExcelDuplication(self,Duplicationexcel='',
                           saveexcel='',
                           header=''):
        excel = pd.read_excel(Duplicationexcel)
        excel.drop_duplicates(subset=header, inplace=True)
        excel.to_excel(saveexcel)

    @pyqtSlot() #�����ظ����ݵ�excel�ļ�
    def on_btn_DuplicationExcel_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEditDuplicationExcel.setText(filename)

    @pyqtSlot() #���������ظ�����ļ�
    def on_btn_SaveDuplicationExcel_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '�����ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getSaveFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_SaveDuplicationExcel.setText(filename)

    @pyqtSlot() #һ������
    def on_btn_ExcelDuplication_clicked(self):
        try:
            self.ui.textEdit_ExcelStatus.clear()
            Duplicationexcel=self.ui.lineEditDuplicationExcel.text()
            saveexcel=self.ui.lineEdit_SaveDuplicationExcel.text()
            header=self.ui.lineEdit_ExcelDuplicationCol.text()
            self.__ExcelDuplication(Duplicationexcel=Duplicationexcel,
                                    saveexcel=saveexcel,header=header)
            self.ui.textEdit_ExcelStatus.setPlainText("�����ɹ���")
            self.ui.lineEditDuplicationExcel.clear()
            self.ui.lineEdit_SaveDuplicationExcel.clear()
            self.ui.lineEdit_ExcelDuplicationCol.clear()
        except:
            self.ui.textEdit_ExcelStatus.clear()
            self.ui.textEdit_ExcelStatus.setPlainText("����ʧ�ܣ������Ƿ������ȷ��")
            self.ui.lineEditDuplicationExcel.clear()
            self.ui.lineEdit_SaveDuplicationExcel.clear()
            self.ui.lineEdit_ExcelDuplicationCol.clear()

#====================�Զ���ת====================
    def __ExcelRotate(self,excelfile='',sheetname='',savefile=''):
        excel = pd.read_excel(excelfile, sheet_name=sheetname, dtype=str)
        table = excel.transpose()
        table.to_excel(savefile)

    @pyqtSlot() #����Ҫ��ת���ļ�
    def on_btn_ExcelRotateFile_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_ExcelRotateFile.setText(filename)

    @pyqtSlot()  # ������ת����ļ�
    def on_btn_SaveExcelRotate_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '�����ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getSaveFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_SaveExcelRotate.setText(filename)

    @pyqtSlot() #һ����ת
    def on_btn_ExcelRotate_clicked(self):
        try:
            self.ui.textEdit_ExcelStatus.clear()
            excelfile=self.ui.lineEdit_ExcelRotateFile.text()
            sheetname=self.ui.lineEdit_ExcelRotatesheet.text()
            savefile=self.ui.lineEdit_SaveExcelRotate.text()
            self.__ExcelRotate(excelfile=excelfile,sheetname=sheetname,savefile=savefile)
            self.ui.textEdit_ExcelStatus.setPlainText("��ת�ɹ���")
            self.ui.lineEdit_ExcelRotateFile.clear()
            self.ui.lineEdit_SaveExcelRotate.clear()
            self.ui.lineEdit_ExcelRotatesheet.clear()
        except:
            self.ui.textEdit_ExcelStatus.clear()
            self.ui.textEdit_ExcelStatus.setPlainText("��תʧ�ܣ���������Ƿ���ȷ��")
            self.ui.lineEdit_ExcelRotateFile.clear()
            self.ui.lineEdit_SaveExcelRotate.clear()
            self.ui.lineEdit_ExcelRotatesheet.clear()




#====================Word���====================
#====================ģ������====================
    #�¾�����ת��
    def change_text(self,document='',old_text='', new_text=''):
        document=Document(document)
        all_paragraphs = document.paragraphs
        for paragraph in all_paragraphs:
            for run in paragraph.runs:
                run_text = run.text.replace(old_text, new_text)
                run.text = run_text
        all_tables = document.tables
        for table in all_tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.replace(old_text, new_text)
                    cell.text = cell_text

    #ģ�����ù���
    def __WordModuleUse(self,wordmodule='',excelmodule='',savelist=''):
        xlsx = xlrd.open_workbook(excelmodule)
        sheet = xlsx.sheet_by_index(0)

        for table_row in range(1, sheet.nrows):
            document = Document(wordmodule)
            for table_col in range(0, sheet.ncols):
                self.change_text(wordmodule, str(sheet.cell_value(0, table_col)), str(sheet.cell_value(table_row, table_col)))

            document.save(savelist + '/' + '%s.docx' % str(sheet.cell_value(table_row, 0)))
            self.ui.textEdit_WordStatus.setText('%s���' % str(sheet.cell_value(table_row, 0)))

    @pyqtSlot() #��Wordģ��
    def on_btn_WordModule_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'Microsoft Word �ĵ�(*.docx)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_WordModule.setText(filename)

    @pyqtSlot()  # ����ģ�����ݵ�excel�ļ�
    def on_btn_WordModuleExcel_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'Microsoft Excel ������(*.xlsx)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_WordModuleExcel.setText(filename)

    @pyqtSlot()  # ѡ��Ҫ����ģ����ļ���
    def on_btn_SaveWordModule_clicked(self):
        curPath = QDir.currentPath()  # ��ȡϵͳ��ǰĿ¼
        dlgTitle = "ѡ��Ŀ¼"
        selectedDir = QFileDialog.getExistingDirectory(self,
                                                       dlgTitle, curPath, QFileDialog.ShowDirsOnly)
        self.ui.lineEdit_SaveWordModule.setText(selectedDir)

    @pyqtSlot()  # ת��Wordģ��
    def on_btn_WordModuleTransform_clicked(self):
        try:
            self.ui.textEdit_WordStatus.clear()
            wordmodule=self.ui.lineEdit_WordModule.text()
            excelmodule=self.ui.lineEdit_WordModuleExcel.text()
            savelist=self.ui.lineEdit_SaveWordModule.text()
            self.__WordModuleUse(wordmodule=wordmodule,excelmodule=excelmodule,savelist=savelist)
            self.ui.textEdit_WordStatus.setPlainText("ת���ɹ���")
            self.ui.lineEdit_WordModule.clear()
            self.ui.lineEdit_WordModuleExcel.clear()
            self.ui.lineEdit_SaveWordModule.clear()
        except:
            self.ui.textEdit_WordStatus.clear()
            self.ui.textEdit_WordStatus.setPlainText("ת��ʧ�ܣ������Ƿ������ȷ��")
            self.ui.lineEdit_WordModule.clear()
            self.ui.lineEdit_WordModuleExcel.clear()
            self.ui.lineEdit_SaveWordModule.clear()

#====================WordתPDF====================
#====================ת��====================
    def __WordToPDF(self,docxpath='',pdfpath=''):
        docx_path = docxpath
        pdf_path = pdfpath

        gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)

        wd = Dispatch("Word.Application")

        doc = wd.Documents.Open(docx_path, ReadOnly=1)
        doc.ExportAsFixedFormat(pdf_path, constants.wdExportFormatPDF, Item=constants.wdExportDocumentWithMarkup,
                                CreateBookmarks=constants.wdExportCreateHeadingBookmarks)

        wd.Quit(constants.wdDoNotSaveChanges)

    @pyqtSlot() #ѡ����Ҫת����word�ļ�
    def on_btn_transWord_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'Microsoft Word �ĵ�(*.docx)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_transWord.setText(filename)

    @pyqtSlot()  # ����ת�����PDF�ļ�
    def on_btn_transPDF_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '�����ļ�'
        filt = 'PDF�ļ�(*.pdf)'
        filename, filtUsed = QFileDialog.getSaveFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_transPDF.setText(filename)

    @pyqtSlot() #һ��ת��
    def on_btn_WordtoPDF_clicked(self):
        try:
            self.ui.textEdit_WordStatus.clear()
            docxpath=self.ui.lineEdit_transWord.text()
            pdfpath=self.ui.lineEdit_transPDF.text()
            self.__WordToPDF(docxpath=docxpath,pdfpath=pdfpath)
            self.ui.textEdit_WordStatus.setPlainText("ת���ɹ���")
            self.ui.lineEdit_transWord.clear()
            self.ui.lineEdit_transPDF.clear()
        except:
            self.ui.textEdit_WordStatus.clear()
            self.ui.textEdit_WordStatus.setPlainText("ת��ʧ�ܣ������Ƿ������ȷ��")
            self.ui.lineEdit_transWord.clear()
            self.ui.lineEdit_transPDF.clear()




#====================PDF���====================
#====================PDF������ȡ====================
    #==========���ֵ�ҳ��ȡ����==========
    def __PDFtextpage(self,filename='',page=1):
        with pdfplumber.open(filename) as pdf:
            text_page=pdf.pages[page-1]
            text=text_page.extract_text()
            self.ui.textEdit_PDF.setPlainText(text)

    #==========��ȡPDFȫ������
    def __PDFtextpages(self,filename=''):
        with pdfplumber.open(filename) as pdf:
            for page in pdf.pages:
                text=page.extract_text()
            self.ui.textEdit_PDF.setPlainText(text)

    @pyqtSlot() #��PDF�ļ�
    def on_btn_PDFTextfile_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = 'PDF �ļ�(*.pdf)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_PDFTextfile.setText(filename)

    @pyqtSlot() #��ȡ��ҳ
    def on_btn_PDFTextgetpage_clicked(self):
        try:
            self.ui.textEdit_PDF.clear()
            self.ui.textEdit_PDFStatus.clear()
            filename=self.ui.lineEdit_PDFTextfile.text()
            page=self.ui.spinBox_PDFTextpage.value()
            self.__PDFtextpage(filename=filename,page=page)
            self.ui.textEdit_PDFStatus.setPlainText("��ȡ�ɹ���")
        except:
            self.ui.textEdit_PDF.clear()
            self.ui.textEdit_PDFStatus.clear()
            self.ui.textEdit_PDFStatus.setPlainText("��ȡʧ�ܣ������Ƿ������ȷ��")

    @pyqtSlot()  # ȫ����ȡ
    def on_btn_PDFTextgetpages_clicked(self):
        try:
            self.ui.textEdit_PDF.clear()
            self.ui.textEdit_PDFStatus.clear()
            filename = self.ui.lineEdit_PDFTextfile.text()
            self.__PDFtextpages(filename=filename)
            self.ui.textEdit_PDFStatus.setPlainText("��ȡ�ɹ���")
        except:
            self.ui.textEdit_PDF.clear()
            self.ui.textEdit_PDFStatus.clear()
            self.ui.textEdit_PDFStatus.setPlainText("��ȡʧ�ܣ������Ƿ������ȷ��")

#====================�ʼ����====================
#====================�ı��ʼ�====================
    #�����ı��ʼ�����
    def __TextEmail(self,host_server='',sender_email='',password='',
                    receiver='',main_title='',main_content=''):
        host_server = host_server  # ����SMTP������
        sender_email = sender_email  # ����������
        password = password  # ��������

        sender_qq_email = sender_email  # ����������
        receiver = receiver  # �ռ�������

        main_title = main_title  # �ʼ�����
        # �ʼ�����
        main_content = main_content

        msg = MIMEMultipart()  # �ʼ�����
        msg['Subject'] = Header(main_title, 'utf-8')
        msg['From'] = sender_qq_email
        msg['To'] = Header('Test', 'utf-8')
        msg.attach(MIMEText(main_content, 'plain', 'utf-8'))

        smtp = SMTP_SSL(host_server)  # ssl��½
        smtp.login(sender_email, password)
        smtp.sendmail(sender_qq_email, receiver, msg.as_string())
        smtp.quit()

    @pyqtSlot() #����ı�����
    def on_btn_TextEmailclearcontent_clicked(self):
        self.ui.textEdit_emailContent.clear()

    @pyqtSlot() #һ������
    def on_btn_TextEmailSend_clicked(self):
        try:
            self.ui.textEdit_emailstatus.clear()
            host_server=self.ui.comboBox_TextEmailsmtpserver.currentText()
            sender_email=self.ui.lineEdit_TextEmailsenderemail.text()
            password=self.ui.lineEdit_TextEmailsenderpassword.text()
            receiver=self.ui.lineEdit_TextEmailreceiveremail.text()
            main_title=self.ui.lineEdit_TextEmailtitle.text()
            main_content=str(self.ui.textEdit_emailContent.toPlainText())
            self.__TextEmail(host_server=host_server,sender_email=sender_email,
                            password=password,receiver=receiver,
                            main_title=main_title,main_content=main_content)
            self.ui.textEdit_emailstatus.setPlainText("���ͳɹ���")
        except:
            self.ui.textEdit_emailstatus.clear()
            self.ui.textEdit_emailstatus.setPlainText("����ʧ�ܣ������Ƿ������ȷ��")

    #�������ʼ�
    def __FileEmail(self,host_server='',sender_email='',password='',
                    receiver='',main_title='',main_content='',filename=''):
        host_server = host_server  # qq����SMTP������
        sender_email = sender_email  # ����������
        password = password  # ������Ȩ��

        sender_qq_email = sender_email  # ����������
        receiver = receiver  # �ռ�������

        main_title = main_title  # �ʼ�����
        # �ʼ�����
        main_content = main_content

        msg = MIMEMultipart()  # �ʼ�����
        msg['Subject'] = Header(main_title, 'utf-8')
        msg['From'] = sender_qq_email
        msg['To'] = Header('Test', 'utf-8')

        msg.attach(MIMEText(main_content, 'plain', 'utf-8'))

        attachment = MIMEApplication(open(filename, 'rb').read())
        new_filename=filename.split("/")[-1]
        attachment.add_header('Content-Disposition', 'attachment', filename=new_filename)

        msg.attach(attachment)

        smtp = SMTP_SSL(host_server)  # ssl��½
        smtp.set_debuglevel(0)
        smtp.ehlo(host_server)
        smtp.login(sender_email, password)
        smtp.sendmail(sender_qq_email, receiver, msg.as_string())
        smtp.quit()

    @pyqtSlot() #����ı�����
    def on_btn_FileEmailclearcontent_clicked(self):
        self.ui.textEdit_emailContent.clear()

    @pyqtSlot() #��ȡ����
    def on_btn_FindFileEmail_clicked(self):
        curPath = QDir.currentPath()
        dlgTitle = '���ļ�'
        filt = '�����ļ�(*.*)'
        filename, filtUsed = QFileDialog.getOpenFileName(self, dlgTitle, curPath, filt)
        self.ui.lineEdit_FindFileEmail.setText(filename)

    @pyqtSlot() #һ������
    def on_btn_FileEmailSend_clicked(self):
        try:
            self.ui.textEdit_emailstatus.clear()
            host_server=self.ui.comboBox_FileEmailsmtpserver.currentText()
            sender_email=self.ui.lineEdit_FileEmailsenderemail.text()
            password=self.ui.lineEdit_FileEmailsenderpassword.text()
            receiver=self.ui.lineEdit_FileEmailreceiveremail.text()
            main_title=self.ui.lineEdit_FileEmailtitle.text()
            main_content=str(self.ui.textEdit_emailContent.toPlainText())
            filename=self.ui.lineEdit_FindFileEmail.text()
            self.__FileEmail(host_server=host_server,sender_email=sender_email,
                            password=password,receiver=receiver,
                            main_title=main_title,main_content=main_content,filename=filename)
            self.ui.textEdit_emailstatus.setPlainText("���ͳɹ���")
        except:
            self.ui.textEdit_emailstatus.clear()
            self.ui.textEdit_emailstatus.setPlainText("����ʧ�ܣ������Ƿ������ȷ��")

#  ============������Գ��� ================================
#if  __name__ == "__main__":        #���ڵ�ǰ�������
#    app = QApplication(sys.argv)    #����GUIӦ�ó���
#    form=QmyMainWindow()            #��������
    #icon = QIcon("logo.ico")
    #app.setWindowIcon(icon)
#    form.show()
#    sys.exit(app.exec_())
